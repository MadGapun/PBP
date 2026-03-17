"""Tests for the PDF/DOCX export module.

Covers CV and cover letter generation in both formats,
verifying files are created, non-empty, and structurally valid.
"""

import pytest
from pathlib import Path


# === Full profile fixture for export tests ===

@pytest.fixture
def full_profile(sample_profile, sample_position, sample_project):
    """Complete profile with positions, education, skills for export testing."""
    profile = dict(sample_profile)
    position = dict(sample_position)
    project = dict(sample_project)
    position["projects"] = [project]
    profile["positions"] = [position]
    profile["education"] = [
        {
            "institution": "TU Hamburg",
            "degree": "Master",
            "field_of_study": "Maschinenbau",
            "start_date": "2010",
            "end_date": "2013",
            "grade": "1.5",
        }
    ]
    profile["skills"] = [
        {"name": "Windchill", "category": "tool", "level": 5},
        {"name": "Python", "category": "tool", "level": 4},
        {"name": "Projektmanagement", "category": "methodisch", "level": 4},
        {"name": "Deutsch", "category": "sprache", "level": 5},
    ]
    return profile


# === CV DOCX ===

class TestCvDocx:
    def test_generate_cv_docx(self, tmp_path, full_profile):
        """Generate CV as DOCX and verify file exists and has content."""
        from bewerbungs_assistent.export import generate_cv_docx
        output = tmp_path / "cv.docx"
        result = generate_cv_docx(full_profile, output)
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0

    def test_cv_docx_content(self, tmp_path, full_profile):
        """Verify DOCX contains expected text sections."""
        from bewerbungs_assistent.export import generate_cv_docx
        from docx import Document
        output = tmp_path / "cv_content.docx"
        generate_cv_docx(full_profile, output)
        doc = Document(str(output))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Max Mustermann" in all_text
        assert "Profil" in all_text
        assert "Berufserfahrung" in all_text or "Tech GmbH" in all_text
        assert "TU Hamburg" in all_text or "Ausbildung" in all_text


# === CV PDF ===

class TestCvPdf:
    def test_generate_cv_pdf(self, tmp_path, full_profile):
        """Generate CV as PDF and verify file exists and is valid PDF."""
        from bewerbungs_assistent.export import generate_cv_pdf
        output = tmp_path / "cv.pdf"
        result = generate_cv_pdf(full_profile, output)
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0
        # Verify it's a real PDF
        with open(output, "rb") as f:
            header = f.read(5)
        assert header == b"%PDF-"

    def test_cv_pdf_minimal_profile(self, tmp_path):
        """CV PDF works with minimal profile (only name)."""
        from bewerbungs_assistent.export import generate_cv_pdf
        output = tmp_path / "cv_minimal.pdf"
        result = generate_cv_pdf({"name": "Test User"}, output)
        assert result == output
        assert output.exists()


# === Cover Letter DOCX ===

class TestCoverLetterDocx:
    def test_generate_cover_letter_docx(self, tmp_path, full_profile):
        """Generate cover letter as DOCX."""
        from bewerbungs_assistent.export import generate_cover_letter_docx
        text = (
            "Sehr geehrte Damen und Herren,\n\n"
            "mit grossem Interesse habe ich Ihre Stellenausschreibung gelesen.\n\n"
            "Ich bringe umfangreiche Erfahrung in PLM-Projekten mit.\n\n"
            "Mit freundlichen Gruessen\nMax Mustermann"
        )
        output = tmp_path / "anschreiben.docx"
        result = generate_cover_letter_docx(
            full_profile, text, "PLM Consultant", "Siemens", output
        )
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0

    def test_cover_letter_docx_content(self, tmp_path, full_profile):
        """Verify cover letter DOCX contains subject line and body."""
        from bewerbungs_assistent.export import generate_cover_letter_docx
        from docx import Document
        text = "Ich bewerbe mich hiermit.\n\nMit freundlichen Gruessen"
        output = tmp_path / "anschreiben_check.docx"
        generate_cover_letter_docx(full_profile, text, "Senior Dev", "TechCo", output)
        doc = Document(str(output))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Bewerbung als Senior Dev" in all_text
        assert "Ich bewerbe mich hiermit" in all_text


# === Cover Letter PDF ===

class TestCoverLetterPdf:
    def test_generate_cover_letter_pdf(self, tmp_path, full_profile):
        """Generate cover letter as PDF."""
        from bewerbungs_assistent.export import generate_cover_letter_pdf
        text = "Sehr geehrte Frau Mueller,\n\nich moechte mich bewerben."
        output = tmp_path / "anschreiben.pdf"
        result = generate_cover_letter_pdf(
            full_profile, text, "Consultant", "Firma XY", output
        )
        assert result == output
        assert output.exists()
        # Verify PDF header
        with open(output, "rb") as f:
            assert f.read(5) == b"%PDF-"

    def test_cover_letter_pdf_empty_text(self, tmp_path, full_profile):
        """Cover letter PDF handles empty text gracefully."""
        from bewerbungs_assistent.export import generate_cover_letter_pdf
        output = tmp_path / "anschreiben_empty.pdf"
        result = generate_cover_letter_pdf(full_profile, "", "Stelle", "Firma", output)
        assert result == output
        assert output.exists()
