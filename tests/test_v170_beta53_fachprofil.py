"""Tests fuer v1.7.0-beta.53 — Kombiniertes Fachprofil-Export (#617)."""
from __future__ import annotations

import os
import tempfile
from pathlib import Path

import pytest


@pytest.fixture
def setup_env():
    tmpdir = tempfile.mkdtemp(prefix="pbp_v170beta53_")
    os.environ["BA_DATA_DIR"] = tmpdir
    import importlib
    import bewerbungs_assistent.database as _db_mod
    importlib.reload(_db_mod)
    import bewerbungs_assistent.dashboard as _dash_mod
    importlib.reload(_dash_mod)
    from bewerbungs_assistent.database import Database
    db = Database()
    db.initialize()
    db.save_profile({
        "name": "Test User",
        "email": "test@example.com",
        "city": "Hamburg",
        "summary": "Senior Engineer mit Fokus auf PLM und Cloud.",
    })
    _dash_mod._db = db
    yield db
    db.close()
    import shutil
    shutil.rmtree(tmpdir, ignore_errors=True)


def _seed_full_profile(db):
    """Profil mit Skills, Positions, Projekten."""
    db.add_skill({"name": "Python", "level": 4, "category": "fachlich"})
    db.add_skill({"name": "PLM", "level": 5, "category": "fachlich"})
    db.add_skill({"name": "Teamcenter", "level": 4, "category": "tool"})
    db.add_skill({"name": "AWS", "level": 3, "category": "tool"})
    db.add_skill({"name": "Englisch", "level": 4, "category": "sprache"})

    pid = db.get_active_profile_id()
    pos1 = db.add_position({
        "title": "Senior PLM Architect", "company": "Beispiel AG",
        "start_date": "2020-01", "end_date": "2024-12",
        "is_current": 0, "description": "Lead-Architekt fuer PLM-Migration",
    })
    pos2 = db.add_position({
        "title": "Tech Lead", "company": "MittelstandTech GmbH",
        "start_date": "2025-01", "end_date": "",
        "is_current": 1, "description": "Technische Leitung Plattform",
    })

    # Projekte unter den Positionen
    db.add_project(pos1, {
        "name": "PLM-Migration Aras",
        "description": "Migration von Aras auf Teamcenter mit 5000 Usern",
        "technologies": "Aras, Teamcenter, Python, REST",
        "result": "Erfolgreich abgeschlossen, 30% Performance-Steigerung",
        "role": "Lead Architect",
        "start_year": 2021, "end_year": 2023,
    })
    db.add_project(pos1, {
        "name": "Schnittstellen-Refactoring",
        "description": "Modernisierung der CAD-Schnittstellen",
        "technologies": "Python, FastAPI, REST",
        "result": "API-Calls von 8s auf 2s reduziert",
        "role": "Senior Engineer",
        "start_year": 2022, "end_year": 2023,
    })
    db.add_project(pos2, {
        "name": "Cloud-Migration",
        "description": "Migration der Plattform nach AWS",
        "technologies": "AWS, Terraform, Docker",
        "result": "Kosten -40%",
        "role": "Tech Lead",
        "start_year": 2025, "end_year": 2025,
    })
    db.add_project(pos1, {
        "name": "Process Excellence Programm",
        "description": "Optimierung interner Workflows",
        "technologies": "Lean, Six Sigma",
        "result": "Durchlaufzeit halbiert",
        "role": "Project Manager",
    })


# ============= export.generate_fachprofil_docx ===============

def test_generate_fachprofil_creates_docx(setup_env, tmp_path):
    db = setup_env
    _seed_full_profile(db)
    profile = db.get_profile()
    from bewerbungs_assistent.export import generate_fachprofil_docx
    out = tmp_path / "fp.docx"
    result = generate_fachprofil_docx(
        profile, "Senior PLM Architect", "Beispiel AG",
        "PLM-Migration mit Teamcenter und Aras-Erfahrung gesucht",
        projekte_anzahl=3, output_path=out,
    )
    assert result == out
    assert out.exists()
    assert out.stat().st_size > 1000  # echte DOCX


def test_fachprofil_docx_contains_target_position(setup_env, tmp_path):
    """Header sollte Zielposition + Firma enthalten."""
    db = setup_env
    _seed_full_profile(db)
    profile = db.get_profile()
    from bewerbungs_assistent.export import generate_fachprofil_docx
    out = tmp_path / "fp.docx"
    generate_fachprofil_docx(
        profile, "Cloud Engineer", "Test-AG",
        projekte_anzahl=2, output_path=out,
    )
    # Inhalt der DOCX lesen
    from docx import Document
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Cloud Engineer" in text
    assert "Test-AG" in text
    assert "Test User" in text  # Profil-Name


def test_fachprofil_picks_top_projects_by_relevance(setup_env, tmp_path):
    """Bei Stellenbeschreibung 'AWS Cloud' sollte Cloud-Migration als
    relevantestes Projekt zuerst kommen."""
    db = setup_env
    _seed_full_profile(db)
    profile = db.get_profile()
    from bewerbungs_assistent.export import generate_fachprofil_docx
    out = tmp_path / "fp.docx"
    generate_fachprofil_docx(
        profile, "Cloud Engineer AWS", "X",
        stellenbeschreibung="AWS Terraform Docker Cloud-Migration",
        projekte_anzahl=2, output_path=out,
    )
    from docx import Document
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    # Cloud-Migration sollte vor Process Excellence kommen
    cloud_pos = text.find("Cloud-Migration")
    process_pos = text.find("Process Excellence")
    assert cloud_pos >= 0
    if process_pos >= 0:
        assert cloud_pos < process_pos


def test_fachprofil_limits_projects(setup_env, tmp_path):
    db = setup_env
    _seed_full_profile(db)  # 4 Projekte gesamt
    profile = db.get_profile()
    from bewerbungs_assistent.export import generate_fachprofil_docx
    out = tmp_path / "fp.docx"
    generate_fachprofil_docx(
        profile, "X", "Y", projekte_anzahl=2, output_path=out,
    )
    from docx import Document
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    # Max 2 Projekte unter "Referenzprojekte"
    # Prozent-Excellence ist niedrigste Priorität — sollte nicht enthalten sein
    proj_section_start = text.find("Referenzprojekte")
    assert proj_section_start >= 0
    section = text[proj_section_start:text.find("Berufliche Stationen", proj_section_start) if "Berufliche Stationen" in text else len(text)]
    proj_count = section.count("\n1.") + section.count("\n2.") + section.count("\n3.")
    # Roughly: 2 enumerated entries
    assert "1." in section and "2." in section
    assert "3." not in section  # cap


def test_fachprofil_works_without_projects(setup_env, tmp_path):
    """Profil ohne Projekte → Sektion 'Referenzprojekte' wird ausgelassen."""
    db = setup_env
    db.add_skill({"name": "Python", "level": 4, "category": "fachlich"})
    profile = db.get_profile()
    from bewerbungs_assistent.export import generate_fachprofil_docx
    out = tmp_path / "fp.docx"
    generate_fachprofil_docx(profile, "X", "Y", projekte_anzahl=5, output_path=out)
    assert out.exists()  # sollte trotzdem funktionieren


def test_fachprofil_pdf_creates_docx_intermediate(setup_env, tmp_path):
    """generate_fachprofil_pdf erzeugt aktuell DOCX als Zwischenstufe."""
    db = setup_env
    _seed_full_profile(db)
    profile = db.get_profile()
    from bewerbungs_assistent.export import generate_fachprofil_pdf
    out = tmp_path / "fp.pdf"
    actual = generate_fachprofil_pdf(profile, "X", "Y", "", 3, out)
    # Aktuell erzeugt PDF-Pfad das DOCX als Zwischenstufe
    assert actual.suffix == ".docx"
    assert actual.exists()


# ============= MCP-Tool fachprofil_exportieren ===============

def _call_tool(db, name, args):
    import asyncio, logging
    from fastmcp import FastMCP
    from bewerbungs_assistent.tools.export_tools import register
    mcp = FastMCP("test")
    register(mcp, db, logging.getLogger("test"))

    async def _run():
        tool = await mcp.get_tool(name)
        res = await tool.run(args)
        return res.structured_content if hasattr(res, "structured_content") else res
    return asyncio.run(_run())


def test_tool_fehler_ohne_profil(setup_env):
    """Wenn kein vollstaendiges Profil → Fehler."""
    # setup_env hat schon ein Profil, aber wir loeschen es absichtlich
    db = setup_env
    pid = db.get_active_profile_id()
    db.delete_profile(pid)
    out = _call_tool(db, "fachprofil_exportieren", {
        "stelle": "Test", "firma": "X",
    })
    assert "fehler" in out


def test_tool_fehler_invalid_format(setup_env):
    db = setup_env
    out = _call_tool(db, "fachprofil_exportieren", {
        "stelle": "X", "firma": "Y", "format": "txt",
    })
    assert "fehler" in out


def test_tool_creates_docx(setup_env):
    db = setup_env
    _seed_full_profile(db)
    out = _call_tool(db, "fachprofil_exportieren", {
        "stelle": "Senior Cloud Engineer",
        "firma": "Test-AG",
        "stellenbeschreibung": "AWS Cloud Migration Engineer",
        "projekte_anzahl": 3,
    })
    assert out["status"] == "erstellt"
    assert out["format"] == "docx"
    assert out["projekte_anzahl_genutzt"] == 3
    assert "fachprofil_" in out["datei"]
    assert Path(out["datei"]).exists()


def test_tool_default_projekte_anzahl_5(setup_env):
    db = setup_env
    _seed_full_profile(db)
    out = _call_tool(db, "fachprofil_exportieren", {
        "stelle": "X", "firma": "Y",
    })
    assert out["projekte_anzahl_genutzt"] == 5


def test_tool_pdf_returns_docx_path(setup_env):
    """PDF-Format erzeugt aktuell DOCX als Zwischenstufe."""
    db = setup_env
    _seed_full_profile(db)
    out = _call_tool(db, "fachprofil_exportieren", {
        "stelle": "X", "firma": "Y", "format": "pdf",
    })
    assert out["status"] == "erstellt"
    # Datei-Endung sollte .docx sein (Zwischenstufe)
    assert out["datei"].endswith(".docx")
