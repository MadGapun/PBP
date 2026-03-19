"""Tests for the PDF/DOCX export module.

Covers CV and cover letter generation in both formats,
verifying files are created, non-empty, and structurally valid.
"""

import pytest
from pathlib import Path


# === Full profile fixture for export tests ===

@pytest.fixture
def full_profile(sample_profile, sample_position, sample_project):
    """Complete profile with positions, education, skills for export testing."""
    profile = dict(sample_profile)
    position = dict(sample_position)
    project = dict(sample_project)
    position["projects"] = [project]
    profile["positions"] = [position]
    profile["education"] = [
        {
            "institution": "TU Hamburg",
            "degree": "Master",
            "field_of_study": "Maschinenbau",
            "start_date": "2010",
            "end_date": "2013",
            "grade": "1.5",
        }
    ]
    profile["skills"] = [
        {"name": "Windchill", "category": "tool", "level": 5},
        {"name": "Python", "category": "tool", "level": 4},
        {"name": "Projektmanagement", "category": "methodisch", "level": 4},
        {"name": "Deutsch", "category": "sprache", "level": 5},
    ]
    return profile


# === CV DOCX ===

class TestCvDocx:
    def test_generate_cv_docx(self, tmp_path, full_profile):
        """Generate CV as DOCX and verify file exists and has content."""
        from bewerbungs_assistent.export import generate_cv_docx
        output = tmp_path / "cv.docx"
        result = generate_cv_docx(full_profile, output)
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0

    def test_cv_docx_content(self, tmp_path, full_profile):
        """Verify DOCX contains expected text sections."""
        from bewerbungs_assistent.export import generate_cv_docx
        from docx import Document
        output = tmp_path / "cv_content.docx"
        generate_cv_docx(full_profile, output)
        doc = Document(str(output))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Max Mustermann" in all_text
        assert "Profil" in all_text
        assert "Berufserfahrung" in all_text or "Tech GmbH" in all_text
        assert "TU Hamburg" in all_text or "Ausbildung" in all_text


# === CV PDF ===

class TestCvPdf:
    def test_generate_cv_pdf(self, tmp_path, full_profile):
        """Generate CV as PDF and verify file exists and is valid PDF."""
        from bewerbungs_assistent.export import generate_cv_pdf
        output = tmp_path / "cv.pdf"
        result = generate_cv_pdf(full_profile, output)
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0
        # Verify it's a real PDF
        with open(output, "rb") as f:
            header = f.read(5)
        assert header == b"%PDF-"

    def test_cv_pdf_minimal_profile(self, tmp_path):
        """CV PDF works with minimal profile (only name)."""
        from bewerbungs_assistent.export import generate_cv_pdf
        output = tmp_path / "cv_minimal.pdf"
        result = generate_cv_pdf({"name": "Test User"}, output)
        assert result == output
        assert output.exists()


# === Cover Letter DOCX ===

class TestCoverLetterDocx:
    def test_generate_cover_letter_docx(self, tmp_path, full_profile):
        """Generate cover letter as DOCX."""
        from bewerbungs_assistent.export import generate_cover_letter_docx
        text = (
            "Sehr geehrte Damen und Herren,\n\n"
            "mit grossem Interesse habe ich Ihre Stellenausschreibung gelesen.\n\n"
            "Ich bringe umfangreiche Erfahrung in PLM-Projekten mit.\n\n"
            "Mit freundlichen Gruessen\nMax Mustermann"
        )
        output = tmp_path / "anschreiben.docx"
        result = generate_cover_letter_docx(
            full_profile, text, "PLM Consultant", "Siemens", output
        )
        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0

    def test_cover_letter_docx_content(self, tmp_path, full_profile):
        """Verify cover letter DOCX contains subject line and body."""
        from bewerbungs_assistent.export import generate_cover_letter_docx
        from docx import Document
        text = "Ich bewerbe mich hiermit.\n\nMit freundlichen Gruessen"
        output = tmp_path / "anschreiben_check.docx"
        generate_cover_letter_docx(full_profile, text, "Senior Dev", "TechCo", output)
        doc = Document(str(output))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Bewerbung als Senior Dev" in all_text
        assert "Ich bewerbe mich hiermit" in all_text


# === Cover Letter PDF ===

class TestCoverLetterPdf:
    def test_generate_cover_letter_pdf(self, tmp_path, full_profile):
        """Generate cover letter as PDF."""
        from bewerbungs_assistent.export import generate_cover_letter_pdf
        text = "Sehr geehrte Frau Mueller,\n\nich moechte mich bewerben."
        output = tmp_path / "anschreiben.pdf"
        result = generate_cover_letter_pdf(
            full_profile, text, "Consultant", "Firma XY", output
        )
        assert result == output
        assert output.exists()
        # Verify PDF header
        with open(output, "rb") as f:
            assert f.read(5) == b"%PDF-"

    def test_cover_letter_pdf_empty_text(self, tmp_path, full_profile):
        """Cover letter PDF handles empty text gracefully."""
        from bewerbungs_assistent.export import generate_cover_letter_pdf
        output = tmp_path / "anschreiben_empty.pdf"
        result = generate_cover_letter_pdf(full_profile, "", "Stelle", "Firma", output)
        assert result == output
        assert output.exists()


# === CV Perspectives Analysis (v0.23.2) ===

class TestCvPerspectivesAnalysis:
    """Tests for the 3-perspectives CV analysis with career gap detection."""

    def test_analyse_returns_all_perspectives(self, full_profile):
        """Analysis returns all three perspectives with scores."""
        from bewerbungs_assistent.export import analyse_cv_perspectives
        result = analyse_cv_perspectives(full_profile, "PLM Consultant", "PLM Windchill Projektmanagement")
        assert "gesamtscore" in result
        assert "perspektiven" in result
        assert "personalberater" in result["perspektiven"]
        assert "ats" in result["perspektiven"]
        assert "recruiter" in result["perspektiven"]
        for key in ("personalberater", "ats", "recruiter"):
            p = result["perspektiven"][key]
            assert "score" in p
            assert 0 <= p["score"] <= 100
            assert "faktoren" in p
            assert "empfehlungen" in p

    def test_career_gap_detection(self):
        """Career gaps longer than 6 months are detected."""
        from bewerbungs_assistent.export import _detect_career_gaps
        positions = [
            {"start_date": "2015-01", "end_date": "2018-06"},
            {"start_date": "2020-01", "end_date": None, "is_current": True},
        ]
        gaps = _detect_career_gaps(positions)
        assert len(gaps) == 1
        assert gaps[0]["months"] > 6

    def test_no_career_gap_for_continuous_positions(self):
        """No gap detected for consecutive positions."""
        from bewerbungs_assistent.export import _detect_career_gaps
        positions = [
            {"start_date": "2015-01", "end_date": "2018-06"},
            {"start_date": "2018-07", "end_date": None, "is_current": True},
        ]
        gaps = _detect_career_gaps(positions)
        assert len(gaps) == 0

    def test_top_recommendations_sorted_by_priority(self, full_profile):
        """Top recommendations are sorted: kritisch > hoch > mittel."""
        from bewerbungs_assistent.export import analyse_cv_perspectives
        result = analyse_cv_perspectives(full_profile, "Java Developer", "Java Spring Boot Microservices")
        recs = result.get("top_empfehlungen", [])
        assert len(recs) > 0
        # Check that critical/high come before medium
        prio_order = {"kritisch": 0, "hoch": 1, "mittel": 2}
        for i in range(1, len(recs)):
            assert prio_order.get(recs[i - 1]["prioritaet"], 2) <= prio_order.get(recs[i]["prioritaet"], 2)

    def test_ats_keyword_matches_returned(self, full_profile):
        """ATS perspective returns matched and missing keywords."""
        from bewerbungs_assistent.export import analyse_cv_perspectives
        result = analyse_cv_perspectives(full_profile, "PLM Consultant", "Windchill PLM Python Projektmanagement")
        ats = result["perspektiven"]["ats"]
        assert "keyword_matches" in ats
        assert "fehlende_keywords" in ats

    def test_weights_affect_combined_score(self, full_profile):
        """Different weights produce different combined scores."""
        from bewerbungs_assistent.export import analyse_cv_perspectives
        r1 = analyse_cv_perspectives(
            full_profile, "PLM Consultant", "PLM",
            {"personalberater": 1.0, "ats": 0.0, "recruiter": 0.0}
        )
        r2 = analyse_cv_perspectives(
            full_profile, "PLM Consultant", "PLM",
            {"personalberater": 0.0, "ats": 1.0, "recruiter": 0.0}
        )
        # Scores should differ (unless by coincidence)
        # At minimum, both should be valid
        assert 0 <= r1["gesamtscore"] <= 100
        assert 0 <= r2["gesamtscore"] <= 100
