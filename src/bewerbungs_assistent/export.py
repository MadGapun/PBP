"""PDF/DOCX Export Module for Bewerbungs-Assistent.

Generates professional CVs and cover letters in PDF and DOCX format.
Uses fpdf2 for PDF (pure Python, no system deps) and python-docx for DOCX.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger("bewerbungs_assistent.export")


def _parse_date_ym(date_str: str) -> Optional[tuple]:
    """Parse a date string to (year, month) tuple for overlap detection."""
    if not date_str:
        return None
    date_str = str(date_str).strip()[:10]
    try:
        if len(date_str) >= 7:
            parts = date_str.split("-")
            return (int(parts[0]), int(parts[1]))
        if len(date_str) == 4:
            return (int(date_str), 1)
    except (ValueError, IndexError):
        pass
    return None


def detect_position_overlaps(positions: list) -> dict:
    """Detect temporal overlaps between positions.

    Returns a dict mapping position index to list of overlapping position indices.
    """
    overlaps = {}
    parsed = []
    now_ym = (datetime.now().year, datetime.now().month)

    for pos in positions:
        start = _parse_date_ym(pos.get("start_date"))
        end = _parse_date_ym(pos.get("end_date"))
        if pos.get("is_current"):
            end = now_ym
        parsed.append((start, end))

    for i in range(len(parsed)):
        s_i, e_i = parsed[i]
        if s_i is None:
            continue
        if e_i is None:
            e_i = now_ym
        for j in range(i + 1, len(parsed)):
            s_j, e_j = parsed[j]
            if s_j is None:
                continue
            if e_j is None:
                e_j = now_ym
            # Check overlap: start_i <= end_j AND start_j <= end_i
            if s_i <= e_j and s_j <= e_i:
                overlaps.setdefault(i, []).append(j)
                overlaps.setdefault(j, []).append(i)

    return overlaps


def _overlap_hint(pos_index: int, positions: list, overlaps: dict) -> str:
    """Generate a hint string for a position with overlaps."""
    if pos_index not in overlaps:
        return ""
    partner_indices = overlaps[pos_index]
    partners = []
    for idx in partner_indices:
        if idx < len(positions):
            name = positions[idx].get("company") or positions[idx].get("title", "")
            partners.append(name)
    if not partners:
        return ""
    return f"(parallel zu {', '.join(partners)})"


def generate_cv_docx(profile: dict, output_path: Path) -> Path:
    """Generate a professional CV as Word document."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Header
    heading = doc.add_heading(profile.get("name", "Lebenslauf"), level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact
    contact = []
    if profile.get("email"): contact.append(profile["email"])
    if profile.get("phone"): contact.append(profile["phone"])
    if profile.get("city"):
        addr = f"{profile.get('address', '')} " if profile.get("address") else ""
        addr += f"{profile.get('plz', '')} {profile['city']}".strip()
        contact.append(addr.strip())
    if contact:
        p = doc.add_paragraph(" | ".join(contact))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary
    if profile.get("summary"):
        doc.add_heading("Profil", level=1)
        doc.add_paragraph(profile["summary"])

    # Work experience
    positions = profile.get("positions", [])
    if positions:
        doc.add_heading("Berufserfahrung", level=1)
        overlaps = detect_position_overlaps(positions)
        for pos_idx, pos in enumerate(positions):
            end = "heute" if pos.get("is_current") else (pos.get("end_date") or "")
            period = f"{pos.get('start_date', '')} - {end}"
            hint = _overlap_hint(pos_idx, positions, overlaps)
            if hint:
                period += f"  {hint}"
            emp_type = pos.get("employment_type", "")
            type_str = f" ({emp_type})" if emp_type and emp_type != "festanstellung" else ""

            p = doc.add_paragraph()
            run = p.add_run(f"{pos.get('title', '')} bei {pos.get('company', '')}{type_str}")
            run.bold = True
            run.font.size = Pt(11)

            p2 = doc.add_paragraph(period)
            if pos.get("location"):
                p2.add_run(f" | {pos['location']}")

            if pos.get("tasks"):
                doc.add_paragraph(f"Aufgaben: {pos['tasks']}")
            if pos.get("achievements"):
                doc.add_paragraph(f"Erfolge: {pos['achievements']}")
            if pos.get("technologies"):
                doc.add_paragraph(f"Technologien: {pos['technologies']}")

            for proj in pos.get("projects", []):
                p = doc.add_paragraph()
                run = p.add_run(f"Projekt: {proj.get('name', '')}")
                run.bold = True
                if proj.get("role"):
                    p.add_run(f" ({proj['role']})")
                if proj.get("result"):
                    doc.add_paragraph(f"Ergebnis: {proj['result']}", style="List Bullet")

    # Education
    education = profile.get("education", [])
    if education:
        doc.add_heading("Ausbildung", level=1)
        for edu in education:
            p = doc.add_paragraph()
            degree = f"{edu.get('degree', '')} {edu.get('field_of_study', '')}".strip()
            run = p.add_run(degree or edu.get("institution", ""))
            run.bold = True
            line = edu.get("institution", "")
            start = edu.get("start_date", "")
            end = edu.get("end_date", "")
            if start or end:
                line += f" | {start} - {end}"
            if edu.get("grade"):
                line += f" | Note: {edu['grade']}"
            doc.add_paragraph(line)

    # Skills
    skills = profile.get("skills", [])
    if skills:
        doc.add_heading("Kompetenzen", level=1)
        by_cat = {}
        for s in skills:
            by_cat.setdefault(s.get("category", "sonstige"), []).append(s)
        cat_labels = {
            "fachlich": "Fachlich", "methodisch": "Methodisch",
            "soft_skill": "Soft Skills", "sprache": "Sprachen",
            "tool": "Tools / Software",
        }
        for cat, items in by_cat.items():
            label = cat_labels.get(cat, cat)
            names = ", ".join(s["name"] for s in items)
            doc.add_paragraph(f"{label}: {names}")

    doc.save(str(output_path))
    logger.info("CV DOCX generated: %s", output_path)
    return output_path


def generate_tailored_cv_docx(
    profile: dict, job_title: str, job_description: str, output_path: Path
) -> Path:
    """Generate a CV tailored for a specific job as Word document.

    Reorders skills and highlights relevant experience based on the job.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Determine relevant keywords from job
    job_text = f"{job_title} {job_description}".lower()
    job_keywords = set(job_text.split())

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Header
    heading = doc.add_heading(profile.get("name", "Lebenslauf"), level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact
    contact = []
    if profile.get("email"): contact.append(profile["email"])
    if profile.get("phone"): contact.append(profile["phone"])
    if profile.get("city"):
        addr = f"{profile.get('address', '')} " if profile.get("address") else ""
        addr += f"{profile.get('plz', '')} {profile['city']}".strip()
        contact.append(addr.strip())
    if contact:
        p = doc.add_paragraph(" | ".join(contact))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Targeted summary — mention the target position
    doc.add_heading("Profil", level=1)
    summary = profile.get("summary", "")
    if summary:
        doc.add_paragraph(summary)

    # Skills — relevant first, grouped by category
    skills = profile.get("skills", [])
    if skills:
        doc.add_heading("Kompetenzen", level=1)

        def skill_relevance(s):
            name_lower = s.get("name", "").lower()
            # Higher relevance if skill name appears in job text
            if name_lower in job_text:
                return 0
            # Check partial match
            if any(kw in name_lower or name_lower in kw for kw in job_keywords if len(kw) > 3):
                return 1
            return 2

        by_cat = {}
        for s in skills:
            by_cat.setdefault(s.get("category", "sonstige"), []).append(s)

        cat_labels = {
            "fachlich": "Fachlich", "methodisch": "Methodisch",
            "soft_skill": "Soft Skills", "sprache": "Sprachen",
            "tool": "Tools / Software",
        }

        # Sort categories: put those with relevant skills first
        def cat_relevance(cat_items):
            cat, items = cat_items
            return min((skill_relevance(s) for s in items), default=2)

        for cat, items in sorted(by_cat.items(), key=cat_relevance):
            label = cat_labels.get(cat, cat)
            # Sort skills within category: relevant first, then by level desc
            sorted_items = sorted(items, key=lambda s: (skill_relevance(s), -(s.get("level", 0) or 0)))
            names = []
            for s in sorted_items:
                name = s["name"]
                level = s.get("level", 0)
                if level and level >= 4:
                    name += " (Experte)" if level == 5 else " (Fortgeschritten)"
                names.append(name)
            doc.add_paragraph(f"{label}: {', '.join(names)}")

    # Work experience — relevant positions first
    positions = profile.get("positions", [])
    if positions:
        doc.add_heading("Berufserfahrung", level=1)

        def pos_relevance(pos):
            pos_text = f"{pos.get('title', '')} {pos.get('tasks', '')} {pos.get('technologies', '')} {pos.get('achievements', '')}".lower()
            hits = sum(1 for kw in job_keywords if len(kw) > 3 and kw in pos_text)
            return -hits  # Negative so most relevant sorts first

        sorted_positions = sorted(positions, key=pos_relevance)
        # Use original positions list for overlap detection (order-independent)
        overlaps = detect_position_overlaps(positions)
        # Map original index for each sorted position
        pos_index_map = {id(p): i for i, p in enumerate(positions)}

        for pos in sorted_positions:
            orig_idx = pos_index_map.get(id(pos), -1)
            end = "heute" if pos.get("is_current") else (pos.get("end_date") or "")
            period = f"{pos.get('start_date', '')} - {end}"
            hint = _overlap_hint(orig_idx, positions, overlaps) if orig_idx >= 0 else ""
            if hint:
                period += f"  {hint}"
            emp_type = pos.get("employment_type", "")
            type_str = f" ({emp_type})" if emp_type and emp_type != "festanstellung" else ""

            p = doc.add_paragraph()
            run = p.add_run(f"{pos.get('title', '')} bei {pos.get('company', '')}{type_str}")
            run.bold = True
            run.font.size = Pt(11)

            p2 = doc.add_paragraph(period)
            if pos.get("location"):
                p2.add_run(f" | {pos['location']}")

            if pos.get("tasks"):
                doc.add_paragraph(f"Aufgaben: {pos['tasks']}")
            if pos.get("achievements"):
                doc.add_paragraph(f"Erfolge: {pos['achievements']}")
            if pos.get("technologies"):
                doc.add_paragraph(f"Technologien: {pos['technologies']}")

            for proj in pos.get("projects", []):
                p = doc.add_paragraph()
                run = p.add_run(f"Projekt: {proj.get('name', '')}")
                run.bold = True
                if proj.get("role"):
                    p.add_run(f" ({proj['role']})")
                if proj.get("result"):
                    doc.add_paragraph(f"Ergebnis: {proj['result']}", style="List Bullet")

    # Education
    education = profile.get("education", [])
    if education:
        doc.add_heading("Ausbildung", level=1)
        for edu in education:
            p = doc.add_paragraph()
            degree = f"{edu.get('degree', '')} {edu.get('field_of_study', '')}".strip()
            run = p.add_run(degree or edu.get("institution", ""))
            run.bold = True
            line = edu.get("institution", "")
            start = edu.get("start_date", "")
            end = edu.get("end_date", "")
            if start or end:
                line += f" | {start} - {end}"
            if edu.get("grade"):
                line += f" | Note: {edu['grade']}"
            doc.add_paragraph(line)

    doc.save(str(output_path))
    logger.info("Tailored CV DOCX generated: %s", output_path)
    return output_path


def analyse_cv_perspectives(
    profile: dict, job_title: str, job_description: str,
    weights: dict = None
) -> dict:
    """Analyse CV-job fit from 3 professional perspectives.

    Perspectives: Personalberater, HR-Recruiter, ATS.
    Each scores 0-100. Combined score uses weights.
    Returns detailed analysis with recommendations.
    """
    import re

    weights = weights or {"personalberater": 0.33, "ats": 0.34, "recruiter": 0.33}
    job_text = f"{job_title} {job_description}".lower()
    job_words = set(w for w in re.findall(r'\b\w+\b', job_text) if len(w) > 2)

    skills = profile.get("skills", [])
    positions = profile.get("positions", [])
    education = profile.get("education", [])

    # --- PERSONALBERATER perspective ---
    pb = _score_personalberater(profile, skills, positions, job_text, job_words)

    # --- ATS perspective ---
    ats = _score_ats(profile, skills, positions, education, job_text, job_words)

    # --- HR-RECRUITER perspective ---
    rec = _score_recruiter(profile, skills, positions, job_text, job_words)

    # Combined score
    combined = round(
        pb["score"] * weights.get("personalberater", 0.33)
        + ats["score"] * weights.get("ats", 0.34)
        + rec["score"] * weights.get("recruiter", 0.33)
    )

    return {
        "gesamtscore": combined,
        "gewichtung": {
            "personalberater": f"{weights.get('personalberater', 0.33) * 100:.0f}%",
            "ats": f"{weights.get('ats', 0.34) * 100:.0f}%",
            "recruiter": f"{weights.get('recruiter', 0.33) * 100:.0f}%",
        },
        "perspektiven": {
            "personalberater": pb,
            "ats": ats,
            "recruiter": rec,
        },
        "top_empfehlungen": _top_recommendations(pb, ats, rec),
    }


def _detect_career_gaps(positions, threshold_months=6):
    """Detect gaps longer than threshold_months between positions.

    Returns list of dicts with 'from' and 'to' date strings.
    """
    import re
    gaps = []
    dated = []
    for pos in positions:
        start = pos.get("start_date", "")
        end = pos.get("end_date", "")
        if not start:
            continue
        # Parse YYYY-MM or YYYY
        m_start = re.match(r"(\d{4})(?:-(\d{1,2}))?", start)
        if not m_start:
            continue
        s_year, s_month = int(m_start.group(1)), int(m_start.group(2) or 1)

        if pos.get("is_current") or not end:
            e_year, e_month = 2026, 3  # current
        else:
            m_end = re.match(r"(\d{4})(?:-(\d{1,2}))?", end)
            if not m_end:
                continue
            e_year, e_month = int(m_end.group(1)), int(m_end.group(2) or 12)

        dated.append({"start": s_year * 12 + s_month, "end": e_year * 12 + e_month,
                       "start_str": start, "end_str": end or "heute"})

    dated.sort(key=lambda x: x["start"])

    for i in range(1, len(dated)):
        gap_months = dated[i]["start"] - dated[i - 1]["end"]
        if gap_months > threshold_months:
            gaps.append({"from": dated[i - 1]["end_str"], "to": dated[i]["start_str"],
                          "months": gap_months})
    return gaps


def _score_personalberater(profile, skills, positions, job_text, job_words):
    """Personalberater perspective: career progression, soft skills, leadership.

    Beruecksichtigt professionelle Recruiter-Kriterien:
    - Lueckenloser, antichronologischer Werdegang
    - Konkrete Erfolge statt Aufgabenbeschreibungen
    - Erkennbare Karriereentwicklung
    - Professionelles Summary/Kurzprofil
    """
    score = 0
    factors = []
    empfehlungen = []

    # Career progression: do titles show growth?
    titles = [p.get("title", "").lower() for p in positions]
    leadership_words = {"lead", "head", "manager", "director", "leiter", "architekt",
                        "principal", "senior", "vp", "chief", "teamleiter", "abteilungsleiter"}
    has_leadership = any(any(lw in t for lw in leadership_words) for t in titles)
    if has_leadership:
        score += 20
        factors.append("Fuehrungserfahrung erkennbar")
    else:
        empfehlungen.append("Fuehrungsverantwortung oder Teamleitung hervorheben (auch informell)")

    # Career length and continuity
    if len(positions) >= 3:
        score += 12
        factors.append(f"{len(positions)} Stationen zeigen breite Erfahrung")
    elif len(positions) >= 1:
        score += 6

    # Career gaps check — Personalberater achten auf lueckenlosen Werdegang
    _gaps = _detect_career_gaps(positions)
    if not _gaps:
        score += 8
        factors.append("Lueckenloser Werdegang")
    else:
        gap_desc = ", ".join(f"{g['from']}-{g['to']}" for g in _gaps[:3])
        empfehlungen.append(
            f"Karriereluecken erkannt ({gap_desc}) — Personalberater fragen danach. "
            "Tipp: Weiterbildungen, Ehrenamt oder Familienzeit eintragen."
        )

    # Soft skills presence
    soft_skills = [s for s in skills if s.get("category") == "soft_skill"]
    if len(soft_skills) >= 3:
        score += 10
        factors.append(f"{len(soft_skills)} Soft Skills dokumentiert")
    elif soft_skills:
        score += 5
        empfehlungen.append("Mehr Soft Skills ergaenzen (Kommunikation, Teamfaehigkeit, Problemloesung)")
    else:
        empfehlungen.append("Soft Skills fehlen komplett — fuer Personalberater ein Warnsignal")

    # Summary/profile statement quality
    summary = profile.get("summary", "")
    if len(summary) > 100:
        score += 12
        factors.append("Aussagekraeftiges Profil-Statement vorhanden")
    elif summary:
        score += 6
        empfehlungen.append("Profil-Statement ausfuehrlicher formulieren (min. 2-3 Saetze)")
    else:
        empfehlungen.append(
            "Profil-Statement fehlt — das ist das Erste was ein Personalberater liest. "
            "Tipp: 3-4 Saetze mit Kernkompetenz, Branchenfokus und Alleinstellungsmerkmal."
        )

    # Projects with results (STAR format)
    projects_with_results = sum(
        1 for p in positions
        for proj in p.get("projects", [])
        if proj.get("result")
    )
    if projects_with_results >= 3:
        score += 15
        factors.append(f"{projects_with_results} Projekte mit messbaren Ergebnissen (STAR)")
    elif projects_with_results >= 1:
        score += 8
        empfehlungen.append("Mehr Projekte mit konkreten Ergebnissen dokumentieren")
    else:
        empfehlungen.append(
            "STAR-Projekte mit messbaren Ergebnissen hinzufuegen — "
            "Personalberater lieben konkrete Erfolge (Budget, Teamgroesse, Zeitersparnis)"
        )

    # Achievements vs. task descriptions — Erfolge statt Aufgabenbeschreibungen
    has_achievements = sum(1 for p in positions if p.get("achievements"))
    has_tasks = sum(1 for p in positions if p.get("tasks"))
    if has_achievements >= len(positions) * 0.6 and positions:
        score += 8
        factors.append("Erfolge bei den meisten Positionen dokumentiert")
    elif has_tasks and not has_achievements:
        empfehlungen.append(
            "Nur Aufgaben, keine Erfolge dokumentiert — "
            "Recruiter wollen wissen WAS du ERREICHT hast, nicht nur was du GETAN hast. "
            "Tipp: Jede Position sollte 1-2 quantifizierte Erfolge haben."
        )

    # Match with job description keywords for cultural/industry fit
    all_text = f"{summary} {' '.join(titles)}".lower()
    industry_hits = sum(1 for w in job_words if w in all_text and len(w) > 4)
    if industry_hits >= 5:
        score += 10
        factors.append("Starke Branchen-/Rollenpassung")
    elif industry_hits >= 2:
        score += 5

    # Contact info completeness (Personalberater erwarten vollstaendige Kontaktdaten)
    contact_complete = all([
        profile.get("name"), profile.get("email"), profile.get("phone"), profile.get("city")
    ])
    if contact_complete:
        score += 5
        factors.append("Kontaktdaten vollstaendig")
    else:
        empfehlungen.append("Kontaktdaten vervollstaendigen (Name, E-Mail, Telefon, Ort)")

    return {
        "score": min(score, 100),
        "label": "Personalberater (Executive Search)",
        "fokus": "Karriereverlauf, Soft Skills, Fuehrung, Branchen-Fit, Lueckenlosigkeit",
        "faktoren": factors,
        "empfehlungen": empfehlungen,
    }


def _score_ats(profile, skills, positions, education, job_text, job_words):
    """ATS perspective: keyword matches, standard format, measurables.

    Beruecksichtigt professionelle Recruiter-Kriterien:
    - Exakte Keyword-Uebereinstimmung (ATS filtert rigoros)
    - Standardisiertes Format und Pflichtfelder
    - Quantifizierte Erfolge mit Zahlen/Metriken
    - Lueckenlose Datumsangaben (Monat/Jahr)
    """
    import re

    score = 0
    factors = []
    empfehlungen = []

    # Exact skill keyword matches
    skill_names = [s.get("name", "").lower() for s in skills]
    exact_matches = []
    for sn in skill_names:
        if sn in job_text:
            exact_matches.append(sn)

    if len(exact_matches) >= 8:
        score += 25
        factors.append(f"{len(exact_matches)} Skills direkt in Stelle gefunden: {', '.join(exact_matches[:8])}")
    elif len(exact_matches) >= 4:
        score += 18
        factors.append(f"{len(exact_matches)} Skill-Treffer: {', '.join(exact_matches)}")
    elif exact_matches:
        score += 8
        factors.append(f"Nur {len(exact_matches)} Skill-Treffer — zu wenig fuer ATS")
        empfehlungen.append("Mehr Keywords aus der Stellenbeschreibung als Skills hinzufuegen")
    else:
        empfehlungen.append(
            "KRITISCH: Kein einziger Skill matcht die Stellenbeschreibung — "
            "ATS wird den CV wahrscheinlich aussortieren. "
            "Tipp: Uebernimm die exakten Begriffe aus der Stellenanzeige."
        )

    # Partial matches (skill word appears in job text)
    partial = []
    for sn in skill_names:
        if sn not in job_text:
            for part in sn.split():
                if len(part) > 3 and part in job_text:
                    partial.append(sn)
                    break
    if partial:
        score += min(len(partial) * 2, 8)
        factors.append(f"{len(partial)} teilweise Treffer")

    # Job title match
    profile_titles = [p.get("title", "").lower() for p in positions]
    title_lower = job_text.split("\n")[0] if job_text else ""
    title_match = any(
        any(tw in pt for tw in re.findall(r'\b\w+\b', title_lower) if len(tw) > 3)
        for pt in profile_titles
    )
    if title_match:
        score += 12
        factors.append("Jobtitel-Uebereinstimmung mit frueheren Positionen")
    else:
        empfehlungen.append(
            "Fruehere Positionstitel aehnlicher zur Zielstelle formulieren — "
            "ATS vergleicht Jobtitel direkt."
        )

    # Measurable achievements (numbers in text)
    all_text = " ".join(
        f"{p.get('achievements', '')} {p.get('tasks', '')}" for p in positions
    )
    numbers = re.findall(
        r'\d+[%\u20ack+]|\d+\s*(?:Prozent|Euro|Mitarbeiter|Projekte|Jahre|Monate|Standorte|Laender)',
        all_text
    )
    if len(numbers) >= 5:
        score += 12
        factors.append(f"{len(numbers)} messbare Erfolge/Metriken gefunden")
    elif len(numbers) >= 2:
        score += 8
        factors.append(f"{len(numbers)} messbare Erfolge gefunden")
        empfehlungen.append("Mehr Zahlen und Metriken einbauen (%, Euro, Teamgroesse, Zeitersparnis)")
    elif numbers:
        score += 4
        empfehlungen.append("Mehr Zahlen und Metriken in Erfolge einbauen (%, Euro, Teamgroesse)")
    else:
        empfehlungen.append(
            "Keine messbaren Erfolge — ATS und Recruiter bevorzugen quantifizierte Ergebnisse. "
            "Tipp: Jede Position sollte mind. eine Zahl enthalten (Budget, Teamgroesse, Ergebnis)."
        )

    # Date format completeness — Monat/Jahr statt nur Jahr
    incomplete_dates = 0
    for pos in positions:
        start = pos.get("start_date", "")
        end = pos.get("end_date", "")
        if start and len(start) <= 4:  # Only year, no month
            incomplete_dates += 1
        if end and len(end) <= 4:
            incomplete_dates += 1
    if incomplete_dates == 0 and positions:
        score += 5
        factors.append("Alle Datumsangaben mit Monat/Jahr")
    elif incomplete_dates > 0:
        empfehlungen.append(
            f"{incomplete_dates} Datumsangaben ohne Monat — "
            "Recruiter erwarten Format MM/JJJJ (z.B. '2019-04')."
        )

    # Education present
    if education:
        score += 8
        factors.append("Ausbildung dokumentiert")
    else:
        empfehlungen.append("Ausbildung fehlt — wird von vielen ATS als Pflichtfeld erwartet")

    # Contact info complete
    has_email = bool(profile.get("email"))
    has_phone = bool(profile.get("phone"))
    has_city = bool(profile.get("city"))
    has_name = bool(profile.get("name"))
    contact_score = sum([has_email, has_phone, has_city, has_name])
    if contact_score == 4:
        score += 8
        factors.append("Kontaktdaten vollstaendig")
    else:
        missing = []
        if not has_name: missing.append("Name")
        if not has_email: missing.append("E-Mail")
        if not has_phone: missing.append("Telefon")
        if not has_city: missing.append("Ort")
        score += contact_score * 2
        empfehlungen.append(f"Kontaktdaten unvollstaendig: {', '.join(missing)} fehlt")

    # Skills with standard categories
    categorized = [s for s in skills if s.get("category") and s["category"] != "sonstige"]
    if len(categorized) >= len(skills) * 0.8 and skills:
        score += 7
        factors.append("Skills gut kategorisiert")
    elif skills:
        empfehlungen.append("Skills besser kategorisieren (fachlich, methodisch, tool, sprache, soft_skill)")

    # Skill levels documented — zeigt Selbsteinschaetzung
    leveled = [s for s in skills if s.get("level") and s["level"] > 0]
    if len(leveled) >= len(skills) * 0.7 and skills:
        score += 5
        factors.append("Skill-Level dokumentiert")

    # Language skills — besonders fuer deutsche ATS relevant
    lang_skills = [s for s in skills if s.get("category") == "sprache"]
    if len(lang_skills) >= 2:
        score += 5
        factors.append(f"{len(lang_skills)} Sprachen dokumentiert")
    elif not lang_skills:
        empfehlungen.append(
            "Sprachkenntnisse fehlen — deutscher Arbeitsmarkt erwartet mindestens Deutsch + Englisch mit Niveau."
        )

    # Compile missing keywords for the user
    all_skill_text = " ".join(skill_names)
    missing_kw = [w for w in job_words if len(w) > 4 and w not in all_skill_text and w.isalpha()]

    return {
        "score": min(score, 100),
        "label": "ATS (Bewerbermanagementsystem)",
        "fokus": "Keyword-Treffer, Standard-Format, messbare Erfolge, vollstaendige Daten, Datumsformat",
        "faktoren": factors,
        "empfehlungen": empfehlungen,
        "keyword_matches": exact_matches,
        "fehlende_keywords": missing_kw[:15],
    }


def _score_recruiter(profile, skills, positions, job_text, job_words):
    """HR-Recruiter perspective: technical depth, project complexity, specific fit.

    Beruecksichtigt professionelle Recruiter-Kriterien:
    - Technische Tiefe statt Breite
    - Branchenrelevante Projektbeschreibungen
    - Erkennbarer roter Faden in der Karriere
    - Konkrete Technologien und Methoden
    """
    score = 0
    factors = []
    empfehlungen = []

    # Technical skills depth (high-level skills)
    expert_skills = [s for s in skills if (s.get("level") or 0) >= 4]
    if len(expert_skills) >= 5:
        score += 20
        names = ", ".join(s["name"] for s in expert_skills[:6])
        factors.append(f"{len(expert_skills)} Expert-Skills (Level 4-5): {names}")
    elif expert_skills:
        score += 12
        empfehlungen.append("Mehr Skills auf Level 4-5 hochstufen wenn die Erfahrung es hergibt")
    else:
        empfehlungen.append("Keine Expert-Level Skills — Recruiter suchen Tiefe, nicht nur Breite")

    # Technology overlap with job
    tool_skills = [s.get("name", "").lower() for s in skills if s.get("category") in ("tool", "fachlich")]
    tech_hits = [t for t in tool_skills if t in job_text]
    if len(tech_hits) >= 5:
        score += 15
        factors.append(f"Starke Tech-Uebereinstimmung: {', '.join(tech_hits[:6])}")
    elif tech_hits:
        score += 8
        factors.append(f"Teilweise Tech-Match: {', '.join(tech_hits)}")
    else:
        empfehlungen.append("Technologie-Stack stimmt kaum ueberein — Skills ggf. ergaenzen")

    # Project complexity (projects with technologies and results)
    complex_projects = sum(
        1 for p in positions
        for proj in p.get("projects", [])
        if proj.get("technologies") or proj.get("result")
    )
    if complex_projects >= 3:
        score += 15
        factors.append(f"{complex_projects} technisch dokumentierte Projekte")
    elif complex_projects >= 1:
        score += 8
    else:
        empfehlungen.append(
            "Projekte mit Technologien und Ergebnissen dokumentieren — "
            "Recruiter in der Fachabteilung wollen Projektdetails sehen."
        )

    # Years of experience (total career span)
    years_exp = 0
    for pos in positions:
        start = pos.get("start_date", "")
        end = pos.get("end_date", "")
        if start:
            try:
                start_year = int(start[:4]) if len(start) >= 4 else 0
                if pos.get("is_current"):
                    end_year = 2026
                elif end and len(end) >= 4:
                    end_year = int(end[:4])
                else:
                    end_year = start_year + 1
                years_exp += max(end_year - start_year, 0)
            except (ValueError, IndexError):
                pass

    if years_exp >= 10:
        score += 12
        factors.append(f"~{years_exp} Jahre Berufserfahrung")
    elif years_exp >= 5:
        score += 8
        factors.append(f"~{years_exp} Jahre Berufserfahrung")
    elif years_exp > 0:
        score += 4

    # Specific tasks/achievements match
    all_tasks = " ".join(
        f"{p.get('tasks', '')} {p.get('achievements', '')} {p.get('technologies', '')}"
        for p in positions
    ).lower()
    task_hits = sum(1 for w in job_words if w in all_tasks and len(w) > 4)
    if task_hits >= 10:
        score += 12
        factors.append(f"Starke Uebereinstimmung in Aufgaben/Erfahrung ({task_hits} Treffer)")
    elif task_hits >= 5:
        score += 7
        factors.append(f"Moderate Aufgaben-Uebereinstimmung ({task_hits} Treffer)")
    else:
        empfehlungen.append("Aufgabenbeschreibungen mit Begriffen aus der Stellenausschreibung anreichern")

    # Red thread — roter Faden: Haben die Positionen einen erkennbaren Zusammenhang?
    all_titles = " ".join(p.get("title", "").lower() for p in positions)
    all_industries = " ".join(
        f"{p.get('company', '')} {p.get('tasks', '')}" for p in positions
    ).lower()
    # Check if job keywords appear across multiple positions
    multi_pos_keywords = 0
    for w in job_words:
        if len(w) > 4:
            positions_mentioning = sum(
                1 for p in positions
                if w in f"{p.get('title', '')} {p.get('tasks', '')} {p.get('technologies', '')}".lower()
            )
            if positions_mentioning >= 2:
                multi_pos_keywords += 1
    if multi_pos_keywords >= 3:
        score += 8
        factors.append("Roter Faden erkennbar — Kernthemen ziehen sich durch mehrere Stationen")
    elif multi_pos_keywords >= 1:
        score += 4

    # Languages
    lang_skills = [s for s in skills if s.get("category") == "sprache"]
    if lang_skills:
        score += 3
        factors.append(f"{len(lang_skills)} Sprache(n) dokumentiert")

    # Certifications / Weiterbildungen (in education or skills)
    certifications = [e for e in (profile.get("education", []) or [])
                      if any(kw in (e.get("degree") or "").lower()
                             for kw in ("zertifik", "certif", "scrum", "prince", "itil", "pmp",
                                        "aws", "azure", "cisco", "sap"))]
    cert_skills = [s for s in skills if s.get("category") == "zertifizierung"]
    total_certs = len(certifications) + len(cert_skills)
    if total_certs >= 2:
        score += 8
        factors.append(f"{total_certs} Zertifizierungen/Weiterbildungen erkannt")
    elif total_certs == 1:
        score += 4
        factors.append("1 Zertifizierung erkannt")
    else:
        empfehlungen.append(
            "Zertifizierungen und Weiterbildungen ergaenzen — "
            "zeigt Engagement und aktuelle Fachkompetenz."
        )

    return {
        "score": min(score, 100),
        "label": "HR-Recruiter (Fachabteilung)",
        "fokus": "Technische Tiefe, Projekt-Komplexitaet, roter Faden, Zertifizierungen",
        "faktoren": factors,
        "empfehlungen": empfehlungen,
    }


def _top_recommendations(pb, ats, rec):
    """Extract top recommendations across all perspectives, prioritized.

    Sortiert nach Wichtigkeit: ATS-Blocker > Recruiter-Concerns > Personalberater-Tipps.
    Markiert kritische Empfehlungen (KRITISCH) als hohe Prioritaet.
    """
    recs = []
    # ATS recommendations first (most critical for getting past the filter)
    for r in ats.get("empfehlungen", []):
        prio = "kritisch" if "KRITISCH" in r else "hoch"
        recs.append({"perspektive": "ATS", "empfehlung": r, "prioritaet": prio})
    for r in rec.get("empfehlungen", []):
        recs.append({"perspektive": "Recruiter", "empfehlung": r, "prioritaet": "mittel"})
    for r in pb.get("empfehlungen", []):
        # Career gaps and missing summary are high priority
        prio = "hoch" if any(kw in r.lower() for kw in ("luecken", "fehlt", "komplett")) else "mittel"
        recs.append({"perspektive": "Personalberater", "empfehlung": r, "prioritaet": prio})

    # Sort: kritisch > hoch > mittel
    prio_order = {"kritisch": 0, "hoch": 1, "mittel": 2}
    recs.sort(key=lambda x: prio_order.get(x["prioritaet"], 2))
    return recs[:8]


def generate_cv_pdf(profile: dict, output_path: Path) -> Path:
    """Generate a professional CV as PDF."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Try Unicode font, fallback to built-in
    try:
        pdf.add_font("DejaVu", "", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
        pdf.add_font("DejaVu", "B", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
        font_name = "DejaVu"
    except Exception as e:
        logger.debug("DejaVu-Font nicht verfuegbar, nutze Helvetica: %s", e)
        font_name = "Helvetica"

    epw = pdf.epw  # effective page width

    def safe(text):
        if not text: return ""
        if font_name == "Helvetica":
            for k, v in {"\u00e4": "ae", "\u00f6": "oe", "\u00fc": "ue",
                         "\u00c4": "Ae", "\u00d6": "Oe", "\u00dc": "Ue",
                         "\u00df": "ss", "\u2013": "-", "\u2014": "-"}.items():
                text = text.replace(k, v)
        return text

    def next_line():
        """Move cursor to left margin on next line."""
        pdf.set_x(pdf.l_margin)

    # Header
    pdf.set_font(font_name, "B", 18)
    pdf.cell(epw, 12, safe(profile.get("name", "Lebenslauf")),
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")

    contact = []
    if profile.get("email"): contact.append(profile["email"])
    if profile.get("phone"): contact.append(profile["phone"])
    if profile.get("city"): contact.append(f"{profile.get('plz', '')} {profile['city']}".strip())
    if contact:
        pdf.set_font(font_name, "", 9)
        pdf.cell(epw, 6, safe(" | ".join(contact)),
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.ln(4)

    def section(title):
        pdf.set_font(font_name, "B", 13)
        pdf.cell(epw, 8, safe(title), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_draw_color(58, 134, 255)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(2)

    # Summary
    if profile.get("summary"):
        section("Profil")
        pdf.set_font(font_name, "", 10)
        next_line()
        pdf.multi_cell(epw, 5, safe(profile["summary"]))
        pdf.ln(3)

    # Positions
    positions = profile.get("positions", [])
    if positions:
        section("Berufserfahrung")
        overlaps = detect_position_overlaps(positions)
        for pos_idx, pos in enumerate(positions):
            end = "heute" if pos.get("is_current") else (pos.get("end_date") or "")
            pdf.set_font(font_name, "B", 11)
            pdf.cell(epw, 6, safe(f"{pos.get('title', '')} bei {pos.get('company', '')}"),
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font(font_name, "", 9)
            loc = f" | {pos['location']}" if pos.get("location") else ""
            hint = _overlap_hint(pos_idx, positions, overlaps)
            hint_str = f"  {hint}" if hint else ""
            pdf.cell(epw, 5, safe(f"{pos.get('start_date', '')} - {end}{loc}{hint_str}"),
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font(font_name, "", 10)
            if pos.get("tasks"):
                next_line()
                pdf.multi_cell(epw, 5, safe(f"Aufgaben: {pos['tasks']}"))
            if pos.get("achievements"):
                next_line()
                pdf.multi_cell(epw, 5, safe(f"Erfolge: {pos['achievements']}"))
            if pos.get("technologies"):
                next_line()
                pdf.multi_cell(epw, 5, safe(f"Technologien: {pos['technologies']}"))
            for proj in pos.get("projects", []):
                pdf.set_font(font_name, "B", 10)
                role = f" ({proj['role']})" if proj.get("role") else ""
                pdf.cell(epw, 5, safe(f"  Projekt: {proj.get('name', '')}{role}"),
                         new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                if proj.get("result"):
                    pdf.set_font(font_name, "", 9)
                    next_line()
                    pdf.multi_cell(epw, 4, safe(f"    Ergebnis: {proj['result']}"))
            pdf.ln(3)

    # Education
    education = profile.get("education", [])
    if education:
        section("Ausbildung")
        for edu in education:
            pdf.set_font(font_name, "B", 10)
            degree = f"{edu.get('degree', '')} {edu.get('field_of_study', '')}".strip()
            pdf.cell(epw, 5, safe(degree or edu.get("institution", "")),
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font(font_name, "", 9)
            line = edu.get("institution", "")
            if edu.get("start_date") or edu.get("end_date"):
                line += f" | {edu.get('start_date', '')} - {edu.get('end_date', '')}"
            if edu.get("grade"): line += f" | Note: {edu['grade']}"
            pdf.cell(epw, 5, safe(line), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(2)

    # Skills
    skills = profile.get("skills", [])
    if skills:
        section("Kompetenzen")
        by_cat = {}
        for s in skills:
            by_cat.setdefault(s.get("category", "sonstige"), []).append(s)
        cat_labels = {"fachlich": "Fachlich", "methodisch": "Methodisch",
                      "soft_skill": "Soft Skills", "sprache": "Sprachen", "tool": "Tools / Software"}
        pdf.set_font(font_name, "", 10)
        for cat, items in by_cat.items():
            label = cat_labels.get(cat, cat)
            names = ", ".join(s["name"] for s in items)
            next_line()
            pdf.multi_cell(epw, 5, safe(f"{label}: {names}"))

    # Footer
    pdf.set_y(-20)
    pdf.set_font(font_name, "", 7)
    pdf.cell(epw, 4, safe(f"Erstellt am {datetime.now().strftime('%d.%m.%Y')}"), align="R")

    pdf.output(str(output_path))
    logger.info("CV PDF generated: %s", output_path)
    return output_path


def generate_cover_letter_docx(
    profile: dict, text: str, stelle: str, firma: str, output_path: Path
) -> Path:
    """Generate a cover letter as Word document."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Sender
    for val in [profile.get("name"), profile.get("address"),
                f"{profile.get('plz', '')} {profile.get('city', '')}".strip(),
                profile.get("phone"), profile.get("email")]:
        if val and val.strip():
            p = doc.add_paragraph(val)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()
    p = doc.add_paragraph(datetime.now().strftime("%d.%m.%Y"))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()

    # Subject
    p = doc.add_paragraph()
    run = p.add_run(f"Bewerbung als {stelle}")
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph()

    # Body
    for paragraph in text.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            doc.add_paragraph(paragraph)

    doc.save(str(output_path))
    logger.info("Cover letter DOCX generated: %s", output_path)
    return output_path


def generate_cover_letter_pdf(
    profile: dict, text: str, stelle: str, firma: str, output_path: Path
) -> Path:
    """Generate a cover letter as PDF."""
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    try:
        pdf.add_font("DejaVu", "", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
        pdf.add_font("DejaVu", "B", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
        font_name = "DejaVu"
    except Exception as e:
        logger.debug("DejaVu-Font nicht verfuegbar, nutze Helvetica: %s", e)
        font_name = "Helvetica"

    epw = pdf.epw

    def safe(t):
        if not t: return ""
        if font_name == "Helvetica":
            for k, v in {"\u00e4": "ae", "\u00f6": "oe", "\u00fc": "ue",
                         "\u00c4": "Ae", "\u00d6": "Oe", "\u00dc": "Ue", "\u00df": "ss"}.items():
                t = t.replace(k, v)
        return t

    # Sender
    pdf.set_font(font_name, "", 10)
    for val in [profile.get("name"), profile.get("address"),
                f"{profile.get('plz', '')} {profile.get('city', '')}".strip(),
                profile.get("phone"), profile.get("email")]:
        if val and val.strip():
            pdf.cell(epw, 5, safe(val), new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="R")
    pdf.ln(8)

    pdf.cell(epw, 5, datetime.now().strftime("%d.%m.%Y"),
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="R")
    pdf.ln(8)

    # Subject
    pdf.set_font(font_name, "B", 12)
    pdf.cell(epw, 7, safe(f"Bewerbung als {stelle}"),
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(6)

    # Body
    pdf.set_font(font_name, "", 11)
    for paragraph in text.split("\n\n"):
        paragraph = paragraph.strip()
        if paragraph:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(epw, 5.5, safe(paragraph))
            pdf.ln(3)

    pdf.output(str(output_path))
    logger.info("Cover letter PDF generated: %s", output_path)
    return output_path


def generate_cv_markdown(profile: dict, output_path: Path) -> Path:
    """Generate a CV as Markdown file."""
    lines = [f"# {profile.get('name', 'Lebenslauf')}", ""]

    contact = []
    if profile.get("email"):
        contact.append(profile["email"])
    if profile.get("phone"):
        contact.append(profile["phone"])
    if profile.get("city"):
        addr = f"{profile.get('address', '')} ".strip()
        addr += f" {profile.get('plz', '')} {profile['city']}".strip()
        contact.append(addr.strip())
    if contact:
        lines.append(" | ".join(contact))
        lines.append("")

    if profile.get("summary"):
        lines.extend(["## Profil", "", profile["summary"], ""])

    positions = profile.get("positions", [])
    if positions:
        lines.append("## Berufserfahrung")
        lines.append("")
        for pos in positions:
            end = "heute" if pos.get("is_current") else (pos.get("end_date") or "")
            period = f"{pos.get('start_date', '')} - {end}"
            emp = pos.get("employment_type", "")
            t = f" ({emp})" if emp and emp != "festanstellung" else ""
            lines.append(f"### {pos.get('title', '')} bei {pos.get('company', '')}{t}")
            lines.append(f"*{period}*")
            lines.append("")
            if pos.get("tasks"):
                lines.append(f"**Aufgaben:** {pos['tasks']}")
            if pos.get("achievements"):
                lines.append(f"**Erfolge:** {pos['achievements']}")
            if pos.get("technologies"):
                lines.append(f"**Technologien:** {pos['technologies']}")
            for proj in pos.get("projects", []):
                lines.append(f"- **Projekt: {proj.get('name', '')}**"
                             + (f" ({proj.get('role', '')})" if proj.get("role") else ""))
                if proj.get("result"):
                    lines.append(f"  Ergebnis: {proj['result']}")
            lines.append("")

    skills = profile.get("skills", [])
    if skills:
        lines.append("## Kompetenzen")
        lines.append("")
        by_cat = {}
        for s in skills:
            cat = s.get("category", "sonstiges")
            by_cat.setdefault(cat, []).append(s.get("name", ""))
        for cat, names in by_cat.items():
            lines.append(f"**{cat}:** {', '.join(names)}")
        lines.append("")

    education = profile.get("education", [])
    if education:
        lines.append("## Ausbildung")
        lines.append("")
        for edu in education:
            period = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}".strip(" -")
            lines.append(f"- **{edu.get('degree', '')}** — {edu.get('institution', '')} ({period})")
        lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")
    logger.info("CV Markdown generated: %s", output_path)
    return output_path


def generate_cv_text(profile: dict, output_path: Path) -> Path:
    """Generate a CV as plain text file."""
    lines = [profile.get("name", "Lebenslauf"), "=" * len(profile.get("name", "Lebenslauf")), ""]

    contact = []
    if profile.get("email"):
        contact.append(profile["email"])
    if profile.get("phone"):
        contact.append(profile["phone"])
    if profile.get("city"):
        addr = f"{profile.get('address', '')} ".strip()
        addr += f" {profile.get('plz', '')} {profile['city']}".strip()
        contact.append(addr.strip())
    if contact:
        lines.append(" | ".join(contact))
        lines.append("")

    if profile.get("summary"):
        lines.extend(["PROFIL", "-" * 6, profile["summary"], ""])

    positions = profile.get("positions", [])
    if positions:
        lines.extend(["BERUFSERFAHRUNG", "-" * 15, ""])
        for pos in positions:
            end = "heute" if pos.get("is_current") else (pos.get("end_date") or "")
            period = f"{pos.get('start_date', '')} - {end}"
            lines.append(f"{pos.get('title', '')} bei {pos.get('company', '')}")
            lines.append(f"  {period}")
            if pos.get("tasks"):
                lines.append(f"  Aufgaben: {pos['tasks']}")
            if pos.get("achievements"):
                lines.append(f"  Erfolge: {pos['achievements']}")
            for proj in pos.get("projects", []):
                lines.append(f"  Projekt: {proj.get('name', '')}"
                             + (f" ({proj.get('role', '')})" if proj.get("role") else ""))
            lines.append("")

    skills = profile.get("skills", [])
    if skills:
        lines.extend(["KOMPETENZEN", "-" * 11, ""])
        by_cat = {}
        for s in skills:
            cat = s.get("category", "sonstiges")
            by_cat.setdefault(cat, []).append(s.get("name", ""))
        for cat, names in by_cat.items():
            lines.append(f"  {cat}: {', '.join(names)}")
        lines.append("")

    education = profile.get("education", [])
    if education:
        lines.extend(["AUSBILDUNG", "-" * 10, ""])
        for edu in education:
            period = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}".strip(" -")
            lines.append(f"  {edu.get('degree', '')} — {edu.get('institution', '')} ({period})")
        lines.append("")

    output_path.write_text("\n".join(lines), encoding="utf-8")
    logger.info("CV text generated: %s", output_path)
    return output_path


def generate_cover_letter_text(
    profile: dict, text: str, stelle: str, firma: str,
    output_path: Path, markdown: bool = False
) -> Path:
    """Generate a cover letter as Markdown or plain text."""
    lines = []
    name = profile.get("name", "")

    if markdown:
        lines.append(f"# Bewerbung als {stelle}")
        lines.append(f"**{firma}**")
        lines.append("")
        if name:
            lines.append(f"*{name}*")
            contact = []
            if profile.get("email"):
                contact.append(profile["email"])
            if profile.get("phone"):
                contact.append(profile["phone"])
            if contact:
                lines.append(f"*{' | '.join(contact)}*")
            lines.append("")
        lines.append(f"*{datetime.now().strftime('%d.%m.%Y')}*")
        lines.append("")
        lines.append("---")
        lines.append("")
        lines.append(text)
    else:
        lines.append(f"Bewerbung als {stelle}")
        lines.append(firma)
        lines.append("")
        if name:
            lines.append(name)
            if profile.get("email"):
                lines.append(profile["email"])
            if profile.get("phone"):
                lines.append(profile["phone"])
            lines.append("")
        lines.append(datetime.now().strftime("%d.%m.%Y"))
        lines.append("")
        lines.append(text)

    output_path.write_text("\n".join(lines), encoding="utf-8")
    logger.info("Cover letter %s generated: %s", "MD" if markdown else "TXT", output_path)
    return output_path
